import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

md_path   = r"C:\claude_workspace\Skattefunn\skattefunn_soknad.md"
docx_path = r"C:\claude_workspace\Skattefunn\skattefunn_soknad.docx"

BLUE = RGBColor(0, 51, 102)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Define styles
def style_heading(para, level):
    para.style = f"Heading {level}"
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.font.color.rgb = BLUE
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 64, 128)
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

def add_body(doc, text):
    # Handle inline bold
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    p.paragraph_format.space_after = Pt(2)

def add_table(doc, rows):
    if not rows: return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = cell_text.strip()
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text.strip())
            run.font.size = Pt(10)
            if ri == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "003366")
                tcPr.append(shd)
    doc.add_paragraph()

def add_code(doc, text):
    text = text.replace("\u2588", "X").replace("\u2592", ".")
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(6)


# --- Parse markdown ---
with open(md_path, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
in_code  = False
code_buf = []
in_table = False
table_rows = []

def flush_table():
    global table_rows, in_table
    if table_rows:
        add_table(doc, table_rows)
    table_rows = []
    in_table = False

while i < len(lines):
    line = lines[i].rstrip("\n")

    if line.strip().startswith("```"):
        if not in_code:
            in_code = True
            code_buf = []
        else:
            add_code(doc, "\n".join(code_buf))
            in_code = False
            code_buf = []
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    if line.strip().startswith("|"):
        in_table = True
        cols = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r'^[-:]+$', c.strip()) for c in cols if c.strip()):
            i += 1
            continue
        table_rows.append(cols)
        i += 1
        continue
    else:
        if in_table:
            flush_table()

    if line.startswith("# "):
        add_h1(doc, line[2:].strip())
    elif line.startswith("## "):
        add_h2(doc, line[3:].strip())
    elif line.startswith("### "):
        add_h3(doc, line[4:].strip())
    elif line.startswith("---"):
        doc.add_paragraph()
    elif line.startswith("- ") or line.startswith("* "):
        add_bullet(doc, line[2:].strip())
    elif line.strip() == "":
        pass
    else:
        clean = re.sub(r'`(.+?)`', r'\1', line)
        add_body(doc, clean)

    i += 1

if in_table:
    flush_table()

doc.save(docx_path)
print(f"Word document created: {docx_path}")
